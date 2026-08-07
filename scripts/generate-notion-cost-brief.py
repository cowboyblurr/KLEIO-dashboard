#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import KeepTogether, LongTable, PageBreak, Paragraph, SimpleDocTemplate, Spacer, TableStyle

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs/notion-attachments/KLEIO_Technical_Cost_Market_Intelligence_Brief_2026-08-07_CURRENT.md"
OUT = ROOT / "docs/notion-attachments"
DOCX = OUT / "KLEIO_Technical_Cost_Market_Intelligence_Brief_2026-08-07_CURRENT.docx"
PDF = OUT / "KLEIO_Technical_Cost_Market_Intelligence_Brief_2026-08-07_CURRENT.pdf"

PURPLE = "5B4B8A"
DEEP = "292631"
MUTED = "6E6879"
LAV = "F4F0FB"
BORDER = "DCD5EC"
LIGHT = "FAF9FC"


def clean_inline(s: str) -> str:
    s = re.sub(r"\*\*(.*?)\*\*", r"\1", s)
    s = re.sub(r"\*(.*?)\*", r"\1", s)
    s = re.sub(r"`(.*?)`", r"\1", s)
    return s.strip()


def rich_html(s: str) -> str:
    # ReportLab paragraph subset.
    s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    s = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", s)
    s = re.sub(r"\*(.*?)\*", r"<i>\1</i>", s)
    s = re.sub(r"`(.*?)`", r"<font name='Courier'>\1</font>", s)
    return s


def parse_blocks(text: str):
    lines = text.splitlines()
    blocks = []
    i = 0
    para = []

    def flush():
        nonlocal para
        if para:
            blocks.append(("p", " ".join(x.strip() for x in para).strip()))
            para = []

    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            flush(); i += 1; continue
        if line.startswith("|"):
            flush()
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [clean_inline(x) for x in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"[-: ]+", x or "-") for x in row):
                    rows.append(row)
                i += 1
            blocks.append(("table", rows)); continue
        if line.startswith("### "):
            flush(); blocks.append(("h3", clean_inline(line[4:]))); i += 1; continue
        if line.startswith("## "):
            flush(); blocks.append(("h2", clean_inline(line[3:]))); i += 1; continue
        if line.startswith("# "):
            flush(); blocks.append(("h1", clean_inline(line[2:]))); i += 1; continue
        if line.startswith("> "):
            flush(); q=[]
            while i < len(lines) and lines[i].startswith("> "):
                q.append(lines[i][2:].strip()); i += 1
            blocks.append(("quote", " ".join(q))); continue
        m = re.match(r"^- (.*)$", line)
        if m:
            flush(); items=[]
            while i < len(lines):
                mm=re.match(r"^- (.*)$", lines[i].rstrip())
                if not mm: break
                items.append(mm.group(1)); i += 1
            blocks.append(("bullets", items)); continue
        m = re.match(r"^\d+\. (.*)$", line)
        if m:
            flush(); items=[]
            while i < len(lines):
                mm=re.match(r"^\d+\. (.*)$", lines[i].rstrip())
                if not mm: break
                items.append(mm.group(1)); i += 1
            blocks.append(("numbered", items)); continue
        para.append(line); i += 1
    flush()
    return blocks


def set_cell_fill(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true"); trPr.append(el)


def set_cant_split(row):
    trPr = row._tr.get_or_add_trPr(); trPr.append(OxmlElement("w:cantSplit"))


def add_docx_table(doc, rows):
    if not rows: return
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(cols):
            val = row[ci] if ci < len(row) else ""
            cells[ci].text = val
            cells[ci].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri == 0:
                set_cell_fill(cells[ci], PURPLE)
                for p in cells[ci].paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255,255,255); run.font.bold=True; run.font.size=Pt(8)
            elif ri % 2 == 0:
                set_cell_fill(cells[ci], LIGHT)
            for p in cells[ci].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs: run.font.size = Pt(8)
        if ri == 0: set_repeat_header(t.rows[-1])
        set_cant_split(t.rows[-1])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build_docx(blocks):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin=Inches(.62); sec.bottom_margin=Inches(.62); sec.left_margin=Inches(.68); sec.right_margin=Inches(.68)
    styles=doc.styles
    styles["Normal"].font.name="Aptos"; styles["Normal"].font.size=Pt(9.4); styles["Normal"].font.color.rgb=RGBColor.from_string(DEEP)
    styles["Normal"].paragraph_format.space_after=Pt(4); styles["Normal"].paragraph_format.line_spacing=1.08
    for n,sz,col in [("Title",24,PURPLE),("Heading 1",15,DEEP),("Heading 2",11.5,PURPLE),("Heading 3",10.3,DEEP)]:
        st=styles[n]; st.font.name="Aptos Display" if n in {"Title","Heading 1"} else "Aptos"; st.font.size=Pt(sz); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(col); st.paragraph_format.keep_with_next=True
    first=True
    for kind,data in blocks:
        if kind=="h1":
            p=doc.add_paragraph(style="Title" if first else "Heading 1"); p.add_run(data); first=False
        elif kind=="h2":
            p=doc.add_paragraph(style="Heading 1"); p.add_run(data)
        elif kind=="h3":
            p=doc.add_paragraph(style="Heading 2"); p.add_run(data)
        elif kind=="p":
            p=doc.add_paragraph(); p.add_run(clean_inline(data))
        elif kind=="quote":
            t=doc.add_table(rows=1, cols=1); c=t.cell(0,0); set_cell_fill(c,LAV); p=c.paragraphs[0]; p.add_run(clean_inline(data)); p.paragraph_format.space_after=Pt(0)
        elif kind=="table": add_docx_table(doc,data)
        elif kind in {"bullets","numbered"}:
            style="List Bullet" if kind=="bullets" else "List Number"
            for item in data:
                p=doc.add_paragraph(style=style); p.paragraph_format.space_after=Pt(2); p.add_run(clean_inline(item))
    for section in doc.sections:
        p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run("KLEIO · Technical Cost & Market Intelligence · Current snapshot · August 7, 2026"); r.font.size=Pt(7.5); r.font.color.rgb=RGBColor.from_string(MUTED)
    doc.save(DOCX)


def build_pdf(blocks):
    styles=getSampleStyleSheet()
    title=ParagraphStyle("TitleK", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=22, leading=25, textColor=colors.HexColor("#"+PURPLE), spaceAfter=8)
    h1=ParagraphStyle("H1K", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=14.5, leading=17, textColor=colors.HexColor("#"+DEEP), spaceBefore=9, spaceAfter=5, keepWithNext=True)
    h2=ParagraphStyle("H2K", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=colors.HexColor("#"+PURPLE), spaceBefore=7, spaceAfter=4, keepWithNext=True)
    body=ParagraphStyle("BodyK", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.7, leading=11.1, textColor=colors.HexColor("#"+DEEP), spaceAfter=4)
    bullet=ParagraphStyle("BulletK", parent=body, leftIndent=12, firstLineIndent=-7, bulletIndent=4, spaceAfter=2)
    quote=ParagraphStyle("QuoteK", parent=body, backColor=colors.HexColor("#"+LAV), borderColor=colors.HexColor("#"+BORDER), borderWidth=.5, borderPadding=7, spaceBefore=4, spaceAfter=7)
    small=ParagraphStyle("SmallK", parent=body, fontSize=7.2, leading=9)
    story=[]; first=True
    for kind,data in blocks:
        if kind=="h1": story.append(Paragraph(rich_html(data), title if first else h1)); first=False
        elif kind=="h2": story.append(Paragraph(rich_html(data), h1))
        elif kind=="h3": story.append(Paragraph(rich_html(data), h2))
        elif kind=="p": story.append(Paragraph(rich_html(data), body))
        elif kind=="quote": story.append(Paragraph(rich_html(data), quote))
        elif kind in {"bullets","numbered"}:
            for idx,item in enumerate(data,1):
                lead="•" if kind=="bullets" else f"{idx}."
                story.append(Paragraph(f"{lead} {rich_html(item)}", bullet))
        elif kind=="table":
            if not data: continue
            pdfrows=[]
            for ri,row in enumerate(data):
                st=small
                pdfrows.append([Paragraph(("<b>"+rich_html(v)+"</b>") if ri==0 else rich_html(v), st) for v in row])
            avail=7.15*inch
            cols=max(len(r) for r in data)
            widths=[avail/cols]*cols
            if cols==2: widths=[avail*.46,avail*.54]
            elif cols==3: widths=[avail*.22,avail*.40,avail*.38]
            elif cols==5: widths=[avail*.34]+[avail*.165]*4
            t=LongTable(pdfrows, colWidths=widths, repeatRows=1, hAlign="LEFT")
            t.setStyle(TableStyle([
                ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#"+PURPLE)),
                ("TEXTCOLOR",(0,0),(-1,0),colors.white),
                ("GRID",(0,0),(-1,-1),.35,colors.HexColor("#"+BORDER)),
                ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
                ("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),
                ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
            ]))
            for r in range(2,len(pdfrows),2): t.setStyle(TableStyle([("BACKGROUND",(0,r),(-1,r),colors.HexColor("#"+LIGHT))]))
            story.extend([t,Spacer(1,5)])
    def footer(canvas, doc):
        canvas.saveState(); canvas.setFillColor(colors.HexColor("#"+MUTED)); canvas.setFont("Helvetica",7)
        canvas.drawCentredString(LETTER[0]/2,0.38*inch,"KLEIO · Technical Cost & Market Intelligence · Current snapshot · August 7, 2026")
        canvas.restoreState()
    pdf=SimpleDocTemplate(str(PDF), pagesize=LETTER, rightMargin=.65*inch,leftMargin=.65*inch,topMargin=.58*inch,bottomMargin=.58*inch, title="KLEIO Technical Cost & Market Intelligence")
    pdf.build(story,onFirstPage=footer,onLaterPages=footer)


def qa():
    forbidden=["Iker","Ortiz","Mister Ortiz","for Iker","Iker's","Iker’s"]
    source=SRC.read_text(encoding="utf-8")
    hits=[x for x in forbidden if x.lower() in source.lower()]
    if hits: raise SystemExit(f"Role-neutral QA failed: {hits}")
    if "earlier August 7 snapshot" in source: raise SystemExit("Stale snapshot warning found")
    for p in (DOCX,PDF):
        if not p.exists() or p.stat().st_size < 10000: raise SystemExit(f"Artifact missing/too small: {p}")
    print(f"Generated {DOCX.relative_to(ROOT)} ({DOCX.stat().st_size} bytes)")
    print(f"Generated {PDF.relative_to(ROOT)} ({PDF.stat().st_size} bytes)")
    print("QA: role-neutral language and current-source gate passed")


if __name__=="__main__":
    blocks=parse_blocks(SRC.read_text(encoding="utf-8"))
    build_docx(blocks)
    build_pdf(blocks)
    qa()
